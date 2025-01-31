import os
import time
from docx import Document

class DocxToMdConverter:
    def __init__(self, docx_path):
        """
        初始化 DocxToMdConverter 类。

        参数：
        docx_path (str): 输入 DOCX 文件路径。
        """
        self.docx_path = docx_path
        self.generate_output_paths()

        # 创建用于保存图片的目录
        os.makedirs(self.image_save_dir, exist_ok=True)

    def generate_output_paths(self):
        """
        生成 Markdown 文件路径和图片保存路径。
        """
        # 获取文件名和路径信息
        base_name = os.path.basename(self.docx_path)  # 获取文件名
        base_dir = os.path.dirname(self.docx_path)  # 获取文件路径
        file_name = os.path.splitext(base_name)[0]  # 文件名去掉扩展名

        # 设置 Markdown 文件保存路径
        self.output_md_path = os.path.join(base_dir.replace("kbs", "kbs_md"), f"{file_name}.md")
        # 设置图片保存路径
        self.image_save_dir = os.path.join(base_dir.replace("kbs", "kbs_imgs"), file_name)

    def save_image(self, blip, document, image_path):
        """
        保存图片到指定路径。

        参数：
        blip: 包含图片数据的 XML 元素。
        document: DOCX 文档对象。
        image_path (str): 图片保存路径。

        返回值：
        bool: 图片保存成功返回 True，否则返回 False。
        """
        # 获取图片的 embed ID 并验证图片是否存在
        embed_id = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
        if embed_id and embed_id in document.part.rels:
            # 获取图片关联信息
            rel = document.part.rels[embed_id]
            if "image" in rel.target_ref:  # 确认是图片资源
                # 保存图片为二进制文件
                with open(image_path, "wb") as img_file:
                    img_file.write(rel.target_part.blob)
                return True
        return False

    def convert(self):
        """
        将 DOCX 文件转换为 Markdown 文件，同时提取图片。
        """
        start_time = time.time()  # 记录处理开始时间
        document = Document(self.docx_path)  # 加载 DOCX 文档
        md_content = []  # 用于存储 Markdown 内容
        image_counter = 0  # 图片计数器

        # 命名空间定义，用于解析 XML 中的图片信息
        namespaces = {
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
            'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        }

        total_paragraphs = len(document.paragraphs)  # 获取文档中的总段落数
        processed_paragraphs = 0  # 已处理的段落数

        # 遍历文档中的每个段落
        for paragraph in document.paragraphs:
            processed_paragraphs += 1
            text = paragraph.text.strip()  # 去除段落的空白字符
            style_name = paragraph.style.name  # 获取段落样式

            # 如果段落为空，跳过处理
            if not text and not paragraph.runs:
                continue

            # 根据段落样式添加 Markdown 标题
            if style_name == "Heading 1":
                md_content.append(f"# {text}\n")  # 一级标题
            elif style_name == "Heading 2":
                md_content.append(f"## {text}\n")  # 二级标题
            elif style_name == "Heading 3":
                md_content.append(f"### {text}\n")  # 三级标题
            else:
                md_content.append(f"{text}\n")  # 普通文本段落

            # 提取段落中的图片
            for run in paragraph.runs:
                for drawing in run.element.findall(".//w:drawing", namespaces):  # 查找图片元素
                    image_counter += 1
                    image_filename = f"image_{image_counter}.png"  # 为图片生成文件名
                    image_path = os.path.join(self.image_save_dir, image_filename)  # 图片保存路径
                    # 保存图片并在 Markdown 文件中插入引用
                    if self.save_image(drawing.find(".//a:blip", namespaces), document, image_path):
                        md_content.append(f"\n![Image {image_counter}]({image_path})\n")

        # 保存生成的 Markdown 文件
        with open(self.output_md_path, "w", encoding="utf-8") as md_file:
            md_file.write("\n".join(md_content))

        # 计算并打印转换所需时间
        elapsed_time = time.time() - start_time
        print(f"转换完成！处理了 {total_paragraphs} 个段落，提取了 {image_counter} 张图片。")
        print(f"耗时：{elapsed_time:.2f} 秒")

        return self.output_md_path  # 返回生成的 Markdown 文件路径


if __name__ == "__main__":
    # 示例用法
    docx_path = "../database/kbs/yyx/测试文档.docx"  # 输入 DOCX 文件路径
    converter = DocxToMdConverter(docx_path)  # 创建转换器实例
    converter.convert()  # 开始转换
